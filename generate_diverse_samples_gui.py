#!/usr/bin/env python3
"""
Diverse Sample Document Generator with GUI
Creates realistic documents in PDF, TXT, and DOCX formats with directory selection.
"""

import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import queue
import random
import io
from datetime import datetime, timedelta
from typing import List, Dict, Tuple

# Import the original generator class
import sys
import importlib.util

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

# DOCX generation
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DiverseDocumentGenerator:
    """Generates realistic documents in multiple formats with proper formatting."""
    
    def __init__(self, output_dir: str = "diverse_sample_documents"):
        self.output_dir = output_dir
        self.setup_data()
        
    def setup_data(self):
        """Setup realistic data for generating documents."""
        self.companies = [
            "TechnoGlobal Solutions Inc", "Meridian Consulting Group", "Apex Digital Systems",
            "Pinnacle Financial Services", "Innovative Research Labs", "Strategic Partners LLC",
            "NextGen Technologies Corp", "Elite Business Solutions", "ProActive Enterprises",
            "Dynamic Consulting Group", "Premier Analytics Inc", "Advanced Systems Ltd"
        ]
        
        self.people = [
            ("John", "Mitchell", "CEO"), ("Sarah", "Chen", "CFO"), ("Michael", "Rodriguez", "CTO"),
            ("Emily", "Thompson", "VP Operations"), ("David", "Kumar", "Legal Counsel"),
            ("Jessica", "Anderson", "HR Director"), ("Robert", "Williams", "Sales Manager"),
            ("Amanda", "Taylor", "Project Manager"), ("James", "Brown", "Senior Analyst"),
            ("Lauren", "Davis", "Marketing Director"), ("Christopher", "Wilson", "Finance Manager"),
            ("Michelle", "Garcia", "Operations Lead")
        ]
        
        self.addresses = [
            ("1250 Tech Park Drive", "San Jose", "CA", "95110"),
            ("890 Business Center Blvd", "Austin", "TX", "78701"),
            ("2100 Corporate Square", "New York", "NY", "10001"),
            ("555 Innovation Way", "Seattle", "WA", "98101"),
            ("1775 Professional Plaza", "Chicago", "IL", "60601"),
            ("3300 Executive Circle", "Dallas", "TX", "75201")
        ]

    def create_output_dir(self):
        """Create output directory structure."""
        os.makedirs(self.output_dir, exist_ok=True)
        for fmt in ['pdf', 'docx', 'txt']:
            os.makedirs(os.path.join(self.output_dir, fmt), exist_ok=True)

    def random_date(self, days_back: int = 365) -> datetime:
        """Generate random date."""
        return datetime.now() - timedelta(days=random.randint(0, days_back))

    def format_currency(self, amount: float) -> str:
        """Format currency amount."""
        return f"${amount:,.2f}"

    # Add all the original generation methods here (abbreviated for space)
    def generate_invoice_pdf(self, filename: str) -> str:
        """Generate a professional invoice PDF."""
        filepath = os.path.join(self.output_dir, 'pdf', filename)
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Simplified invoice generation
        company = random.choice(self.companies)
        invoice_num = f"INV-{random.randint(2023, 2024)}-{random.randint(1000, 9999)}"
        date = self.random_date(90)
        
        story.append(Paragraph(f"<b>{company}</b>", styles['Title']))
        story.append(Paragraph("INVOICE", styles['Heading1']))
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"Invoice #: {invoice_num}", styles['Normal']))
        story.append(Paragraph(f"Date: {date.strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Services and total
        story.append(Paragraph("Professional Services: $5,000.00", styles['Normal']))
        story.append(Paragraph("Tax: $437.50", styles['Normal']))
        story.append(Paragraph("<b>Total Due: $5,437.50</b>", styles['Normal']))
        
        doc.build(story)
        return filepath

    def generate_memo_docx(self, filename: str) -> str:
        """Generate a professional memo DOCX."""
        filepath = os.path.join(self.output_dir, 'docx', filename)
        doc = DocxDocument()
        
        sender_name, sender_last, sender_title = random.choice(self.people)
        memo_date = self.random_date(30)
        subject = "Important Company Update"
        
        doc.add_heading('MEMORANDUM', 0)
        doc.add_paragraph(f"TO: All Staff")
        doc.add_paragraph(f"FROM: {sender_name} {sender_last}, {sender_title}")
        doc.add_paragraph(f"DATE: {memo_date.strftime('%B %d, %Y')}")
        doc.add_paragraph(f"RE: {subject}")
        doc.add_paragraph("")
        doc.add_paragraph("This memo provides important updates regarding company policies and procedures.")
        
        doc.save(filepath)
        return filepath

    def generate_contract_txt(self, filename: str) -> str:
        """Generate a contract TXT file."""
        filepath = os.path.join(self.output_dir, 'txt', filename)
        company1 = random.choice(self.companies)
        company2 = random.choice([c for c in self.companies if c != company1])
        date = self.random_date(180)
        
        content = f"""SERVICE AGREEMENT

This Service Agreement ("Agreement") is entered into on {date.strftime('%B %d, %Y')}
between {company1} ("Provider") and {company2} ("Client").

TERMS AND CONDITIONS:

1. SCOPE OF SERVICES
Provider agrees to deliver professional consulting services as outlined in Exhibit A.

2. PAYMENT TERMS
Client agrees to pay Provider the total amount of ${random.randint(10000, 50000):,}.00
within 30 days of invoice receipt.

3. DURATION
This agreement shall remain in effect for a period of 12 months from the effective date.

4. TERMINATION
Either party may terminate this agreement with 30 days written notice.

By signing below, both parties agree to the terms and conditions outlined above.

_________________________        _________________________
{company1}                         {company2}
Provider                          Client

Date: _______________            Date: _______________
"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return filepath

    def generate_legal_pdf(self, filename: str) -> str:
        """Generate a legal document PDF."""
        filepath = os.path.join(self.output_dir, 'pdf', filename)
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        case_num = f"Case No. {random.randint(2023, 2024)}-{random.randint(100, 999)}"
        date = self.random_date(60)
        
        story.append(Paragraph("LEGAL NOTICE", styles['Title']))
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"{case_num}", styles['Heading2']))
        story.append(Paragraph(f"Date: {date.strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 20))
        story.append(Paragraph("This document serves as official legal notice regarding the matter at hand.", styles['Normal']))
        story.append(Paragraph("All parties are hereby notified of their rights and obligations under applicable law.", styles['Normal']))
        
        doc.build(story)
        return filepath

    def generate_report_docx(self, filename: str) -> str:
        """Generate a report DOCX file."""
        filepath = os.path.join(self.output_dir, 'docx', filename)
        doc = DocxDocument()
        
        company = random.choice(self.companies)
        author_name, author_last, author_title = random.choice(self.people)
        date = self.random_date(30)
        
        doc.add_heading('QUARTERLY BUSINESS REPORT', 0)
        doc.add_paragraph(f"Company: {company}")
        doc.add_paragraph(f"Prepared by: {author_name} {author_last}, {author_title}")
        doc.add_paragraph(f"Date: {date.strftime('%B %d, %Y')}")
        doc.add_paragraph("")
        
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph("This quarterly report provides an overview of business performance and key metrics.")
        
        doc.add_heading('Financial Performance', level=1)
        doc.add_paragraph(f"Revenue: ${random.randint(500000, 2000000):,}")
        doc.add_paragraph(f"Expenses: ${random.randint(300000, 1500000):,}")
        doc.add_paragraph(f"Net Income: ${random.randint(50000, 500000):,}")
        
        doc.save(filepath)
        return filepath

    def generate_other_txt(self, filename: str) -> str:
        """Generate other document types as TXT."""
        filepath = os.path.join(self.output_dir, 'txt', filename)
        
        doc_types = ["Technical Manual", "User Guide", "Reference Document", "Meeting Minutes"]
        doc_type = random.choice(doc_types)
        date = self.random_date(60)
        
        content = f"""{doc_type.upper()}

Document Type: {doc_type}
Generated: {date.strftime('%B %d, %Y')}
Version: 1.0

OVERVIEW
This document provides technical information and guidelines for system users.

SECTIONS
1. Introduction and Overview
2. System Requirements
3. Installation Procedures
4. Configuration Guidelines
5. Troubleshooting Information

NOTES
Please refer to the latest documentation for current procedures and requirements.
Contact technical support for additional assistance.
"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return filepath

    def generate_all_diverse_documents(self, total_docs: int = 60, progress_callback=None) -> Dict[str, int]:
        """Generate diverse documents across all formats and categories."""
        self.create_output_dir()
        
        # Define distribution
        docs_per_category = total_docs // 6
        remainder = total_docs % 6
        
        distribution = {
            'invoice': docs_per_category + (1 if remainder > 0 else 0),
            'memo': docs_per_category + (1 if remainder > 1 else 0),
            'contract': docs_per_category + (1 if remainder > 2 else 0),
            'legal': docs_per_category + (1 if remainder > 3 else 0),
            'report': docs_per_category + (1 if remainder > 4 else 0),
            'other': docs_per_category + (1 if remainder > 5 else 0)
        }
        
        generated_files = []
        format_count = {'pdf': 0, 'docx': 0, 'txt': 0}
        total_generated = 0
        
        # Generate each category
        categories = [
            ('invoice', ['pdf', 'txt'], [0.8, 0.2]),
            ('memo', ['docx', 'txt'], [0.7, 0.3]),
            ('contract', ['txt', 'pdf', 'docx'], [0.6, 0.2, 0.2]),
            ('legal', ['pdf', 'txt'], [0.8, 0.2]),
            ('report', ['docx', 'pdf'], [0.7, 0.3]),
            ('other', ['txt', 'pdf', 'docx'], [0.4, 0.3, 0.3])
        ]
        
        for category, formats, weights in categories:
            for i in range(distribution[category]):
                fmt = random.choices(formats, weights=weights)[0]
                filename = f"{category}_{i+1:03d}.{fmt}"
                
                # Generate the document
                try:
                    if category == 'invoice':
                        if fmt == 'pdf':
                            self.generate_invoice_pdf(filename)
                        else:
                            # Create simplified invoice txt
                            self.generate_simple_txt(filename, 'INVOICE', 'Invoice documentation')
                    elif category == 'memo':
                        if fmt == 'docx':
                            self.generate_memo_docx(filename)
                        else:
                            self.generate_simple_txt(filename, 'MEMO', 'Internal memorandum')
                    elif category == 'contract':
                        if fmt == 'txt':
                            self.generate_contract_txt(filename)
                        elif fmt == 'pdf':
                            self.generate_simple_pdf(filename, 'CONTRACT', 'Service agreement document')
                        else:
                            self.generate_simple_docx(filename, 'CONTRACT', 'Legal contract document')
                    elif category == 'legal':
                        if fmt == 'pdf':
                            self.generate_legal_pdf(filename)
                        else:
                            self.generate_simple_txt(filename, 'LEGAL DOCUMENT', 'Legal notice and documentation')
                    elif category == 'report':
                        if fmt == 'docx':
                            self.generate_report_docx(filename)
                        else:
                            self.generate_simple_pdf(filename, 'BUSINESS REPORT', 'Quarterly business analysis')
                    elif category == 'other':
                        if fmt == 'txt':
                            self.generate_other_txt(filename)
                        elif fmt == 'pdf':
                            self.generate_simple_pdf(filename, 'TECHNICAL DOCUMENT', 'Technical reference material')
                        else:
                            self.generate_simple_docx(filename, 'REFERENCE GUIDE', 'Technical documentation')
                    
                    format_count[fmt] += 1
                    generated_files.append(filename)
                    total_generated += 1
                    
                    # Progress callback
                    if progress_callback:
                        progress = (total_generated / total_docs) * 100
                        progress_callback(progress, f"Generated {filename}")
                        
                except Exception as e:
                    if progress_callback:
                        progress_callback(None, f"Error generating {filename}: {str(e)}")
        
        return format_count

    def generate_simple_txt(self, filename: str, doc_type: str, description: str):
        """Generate a simple TXT document."""
        filepath = os.path.join(self.output_dir, 'txt', filename)
        date = self.random_date(90)
        company = random.choice(self.companies)
        
        content = f"""{doc_type}

Company: {company}
Date: {date.strftime('%B %d, %Y')}
Document ID: {random.randint(1000, 9999)}

{description}

This document contains important information regarding business operations and procedures.
Please review all sections carefully and contact the appropriate department for questions.

Generated on {datetime.now().strftime('%Y-%m-%d')}
"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

    def generate_simple_pdf(self, filename: str, doc_type: str, description: str):
        """Generate a simple PDF document."""
        filepath = os.path.join(self.output_dir, 'pdf', filename)
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        date = self.random_date(90)
        company = random.choice(self.companies)
        
        story.append(Paragraph(doc_type, styles['Title']))
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"Company: {company}", styles['Normal']))
        story.append(Paragraph(f"Date: {date.strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 20))
        story.append(Paragraph(description, styles['Normal']))
        story.append(Paragraph("This document contains important business information.", styles['Normal']))
        
        doc.build(story)

    def generate_simple_docx(self, filename: str, doc_type: str, description: str):
        """Generate a simple DOCX document."""
        filepath = os.path.join(self.output_dir, 'docx', filename)
        doc = DocxDocument()
        
        date = self.random_date(90)
        company = random.choice(self.companies)
        
        doc.add_heading(doc_type, 0)
        doc.add_paragraph(f"Company: {company}")
        doc.add_paragraph(f"Date: {date.strftime('%B %d, %Y')}")
        doc.add_paragraph("")
        doc.add_paragraph(description)
        doc.add_paragraph("This document contains important business information and procedures.")
        
        doc.save(filepath)


class DocumentGeneratorGUI:
    """GUI application for the diverse document generator."""
    
    def __init__(self, root):
        self.root = root
        self.root.title("📄 Diverse Document Generator")
        self.root.geometry("700x600")
        
        # Variables
        self.output_dir = tk.StringVar(value=os.path.join(os.getcwd(), "diverse_sample_documents"))
        self.num_docs = tk.StringVar(value="60")
        self.is_generating = False
        
        # Queue for thread communication
        self.queue = queue.Queue()
        
        self.create_widgets()
        self.check_queue()
        
    def create_widgets(self):
        """Create and arrange GUI widgets."""
        
        # Main frame
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Title
        title_label = ttk.Label(main_frame, text="📄 Diverse Document Generator", 
                               font=('Arial', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # Description
        desc_label = ttk.Label(main_frame, 
                              text="Generate realistic PDF, DOCX, and TXT documents for testing",
                              font=('Arial', 10))
        desc_label.grid(row=1, column=0, columnspan=3, pady=(0, 20))
        
        # Output directory selection
        ttk.Label(main_frame, text="Output Directory:", font=('Arial', 10, 'bold')).grid(
            row=2, column=0, sticky=tk.W, pady=(0, 5))
        
        dir_frame = ttk.Frame(main_frame)
        dir_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 15))
        
        self.dir_entry = ttk.Entry(dir_frame, textvariable=self.output_dir, width=50)
        self.dir_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        
        ttk.Button(dir_frame, text="Browse", command=self.browse_directory).grid(
            row=0, column=1)
        
        dir_frame.columnconfigure(0, weight=1)
        
        # Number of documents
        ttk.Label(main_frame, text="Number of Documents:", font=('Arial', 10, 'bold')).grid(
            row=4, column=0, sticky=tk.W, pady=(0, 5))
        
        num_frame = ttk.Frame(main_frame)
        num_frame.grid(row=5, column=0, columnspan=3, sticky=tk.W, pady=(0, 15))
        
        ttk.Entry(num_frame, textvariable=self.num_docs, width=10).grid(row=0, column=0)
        ttk.Label(num_frame, text="(Recommended: 50-1000)").grid(row=0, column=1, padx=(10, 0))
        
        # Document type distribution info
        info_frame = ttk.LabelFrame(main_frame, text="Document Distribution", padding="10")
        info_frame.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 15))
        
        distribution_text = """• Invoices: 80% PDF, 20% TXT
• Memos: 70% DOCX, 30% TXT  
• Contracts: 60% TXT, 20% PDF, 20% DOCX
• Legal: 80% PDF, 20% TXT
• Reports: 70% DOCX, 30% PDF
• Other: 40% TXT, 30% PDF, 30% DOCX"""
        
        ttk.Label(info_frame, text=distribution_text).grid(row=0, column=0, sticky=tk.W)
        
        # Progress section
        progress_frame = ttk.LabelFrame(main_frame, text="Progress", padding="10")
        progress_frame.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 15))
        
        self.progress_var = tk.StringVar(value="Ready to generate documents")
        ttk.Label(progress_frame, textvariable=self.progress_var).grid(row=0, column=0, sticky=tk.W)
        
        self.progress_bar = ttk.Progressbar(progress_frame, length=400, mode='determinate')
        self.progress_bar.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(5, 0))
        
        progress_frame.columnconfigure(0, weight=1)
        
        # Output log
        log_frame = ttk.LabelFrame(main_frame, text="Output Log", padding="10")
        log_frame.grid(row=8, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 15))
        
        self.log_text = scrolledtext.ScrolledText(log_frame, width=70, height=10, state='disabled')
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=9, column=0, columnspan=3, pady=(10, 0))
        
        self.generate_btn = ttk.Button(button_frame, text="🚀 Generate Documents", 
                                      command=self.start_generation, style='Accent.TButton')
        self.generate_btn.grid(row=0, column=0, padx=(0, 10))
        
        ttk.Button(button_frame, text="📂 Open Output Folder", 
                  command=self.open_output_folder).grid(row=0, column=1, padx=(0, 10))
        
        ttk.Button(button_frame, text="🧹 Clear Log", 
                  command=self.clear_log).grid(row=0, column=2)
        
        # Configure grid weights
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(8, weight=1)
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
    def browse_directory(self):
        """Open directory browser."""
        directory = filedialog.askdirectory(initialdir=self.output_dir.get())
        if directory:
            self.output_dir.set(directory)
            
    def log_message(self, message):
        """Add message to log."""
        self.log_text.config(state='normal')
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)
        self.log_text.config(state='disabled')
        
    def clear_log(self):
        """Clear the log text."""
        self.log_text.config(state='normal')
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state='disabled')
        
    def open_output_folder(self):
        """Open the output folder in file explorer."""
        output_path = self.output_dir.get()
        if os.path.exists(output_path):
            os.startfile(output_path)
        else:
            messagebox.showwarning("Warning", "Output directory does not exist yet!")
            
    def progress_callback(self, progress, message):
        """Callback for progress updates from generator thread."""
        self.queue.put(('progress', progress, message))
        
    def start_generation(self):
        """Start document generation in a separate thread."""
        if self.is_generating:
            return
            
        try:
            num_docs = int(self.num_docs.get())
            if num_docs <= 0 or num_docs > 10000:
                messagebox.showerror("Error", "Please enter a number between 1 and 10000")
                return
        except ValueError:
            messagebox.showerror("Error", "Please enter a valid number")
            return
            
        output_dir = self.output_dir.get()
        if not output_dir:
            messagebox.showerror("Error", "Please select an output directory")
            return
            
        self.is_generating = True
        self.generate_btn.config(state='disabled', text="⏳ Generating...")
        self.progress_bar['value'] = 0
        self.log_message(f"Starting generation of {num_docs} documents...")
        self.log_message(f"Output directory: {output_dir}")
        
        # Start generation in separate thread
        thread = threading.Thread(target=self.generate_documents, 
                                 args=(output_dir, num_docs))
        thread.daemon = True
        thread.start()
        
    def generate_documents(self, output_dir, num_docs):
        """Generate documents in separate thread."""
        try:
            generator = DiverseDocumentGenerator(output_dir)
            format_counts = generator.generate_all_diverse_documents(
                num_docs, self.progress_callback)
            
            self.queue.put(('complete', format_counts))
            
        except Exception as e:
            self.queue.put(('error', str(e)))
            
    def check_queue(self):
        """Check for messages from generator thread."""
        try:
            while True:
                msg = self.queue.get_nowait()
                
                if msg[0] == 'progress':
                    _, progress, message = msg
                    if progress is not None:
                        self.progress_bar['value'] = progress
                        self.progress_var.set(f"Progress: {progress:.1f}%")
                    self.log_message(message)
                    
                elif msg[0] == 'complete':
                    _, format_counts = msg
                    self.generation_complete(format_counts)
                    
                elif msg[0] == 'error':
                    _, error_msg = msg
                    self.generation_error(error_msg)
                    
        except queue.Empty:
            pass
            
        # Check again in 100ms
        self.root.after(100, self.check_queue)
        
    def generation_complete(self, format_counts):
        """Handle completion of document generation."""
        self.is_generating = False
        self.generate_btn.config(state='normal', text="🚀 Generate Documents")
        self.progress_bar['value'] = 100
        self.progress_var.set("Generation complete!")
        
        total_docs = sum(format_counts.values())
        self.log_message(f"✅ Generation complete! Created {total_docs} documents:")
        self.log_message(f"   📄 PDFs: {format_counts.get('pdf', 0)}")
        self.log_message(f"   📝 DOCX: {format_counts.get('docx', 0)}")
        self.log_message(f"   📋 TXT: {format_counts.get('txt', 0)}")
        self.log_message(f"📂 Files saved to: {self.output_dir.get()}")
        
        messagebox.showinfo("Success", 
                           f"Successfully generated {total_docs} documents!\n\n"
                           f"Location: {self.output_dir.get()}")
        
    def generation_error(self, error_msg):
        """Handle generation error."""
        self.is_generating = False
        self.generate_btn.config(state='normal', text="🚀 Generate Documents")
        self.progress_var.set("Generation failed")
        
        self.log_message(f"❌ Error: {error_msg}")
        messagebox.showerror("Error", f"Generation failed:\n{error_msg}")


def main():
    """Main function to launch the GUI."""
    root = tk.Tk()
    
    # Configure style
    style = ttk.Style()
    style.theme_use('clam')
    
    app = DocumentGeneratorGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main() 